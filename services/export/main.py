"""Export Service - Multi-format export system for web scraping results."""

import asyncio
import csv
import json
import os
import time
import uuid
from datetime import datetime
from enum import Enum
from io import BytesIO, StringIO
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

import structlog
from docx import Document
from docx.shared import Inches
from fastapi import BackgroundTasks, FastAPI, HTTPException, Query
from jinja2 import Environment, FileSystemLoader, Template
from notion_client import Client as NotionClient
from pydantic import BaseModel, Field
from qdrant_client import QdrantClient
from qdrant_client.http import models

structlog.configure(
    processors=[
        structlog.stdlib.add_log_level,
        structlog.processors.TimeStamper(fmt="iso"),
        structlog.processors.JSONRenderer()
    ],
    logger_factory=structlog.stdlib.LoggerFactory(),
    wrapper_class=structlog.stdlib.BoundLogger,
    cache_logger_on_first_use=True,
)

logger = structlog.get_logger()

app = FastAPI(
    title="Export Service",
    description="Multi-format export system for web scraping and clustering results",
    version="1.0.0"
)

# Configuration
QDRANT_HOST = os.getenv("QDRANT_HOST", "qdrant")
QDRANT_PORT = int(os.getenv("QDRANT_PORT", "6333"))
EXPORT_DIR = Path(os.getenv("EXPORT_DIR", "/app/exports"))
TEMPLATES_DIR = Path(os.getenv("TEMPLATES_DIR", "/app/templates"))

# For local testing, use relative paths
if not EXPORT_DIR.exists():
    EXPORT_DIR = Path("./exports")
if not TEMPLATES_DIR.exists():
    TEMPLATES_DIR = Path("./templates")

# Ensure directories exist
EXPORT_DIR.mkdir(exist_ok=True)
TEMPLATES_DIR.mkdir(exist_ok=True)

# Initialize Qdrant client
qdrant_client = QdrantClient(host=QDRANT_HOST, port=QDRANT_PORT)

# Export formats
class ExportFormat(str, Enum):
    JSON = "json"
    CSV = "csv"
    MARKDOWN = "markdown"
    WORD = "word"
    NOTION = "notion"
    OBSIDIAN = "obsidian"
    PDF = "pdf"

class ExportStatus(str, Enum):
    PENDING = "pending"
    PROCESSING = "processing"
    COMPLETED = "completed"
    FAILED = "failed"

# Data models
class ExportFilter(BaseModel):
    cluster_ids: Optional[List[int]] = None
    date_range: Optional[Dict[str, str]] = None
    content_types: Optional[List[str]] = None
    domains: Optional[List[str]] = None
    min_score: Optional[float] = None
    keywords: Optional[List[str]] = None

class ExportTemplate(BaseModel):
    name: str
    format: ExportFormat
    template_content: str
    variables: Dict[str, Any] = Field(default_factory=dict)

class ExportRequest(BaseModel):
    session_id: str
    format: ExportFormat
    template_name: Optional[str] = None
    custom_template: Optional[str] = None
    filters: Optional[ExportFilter] = None
    include_metadata: bool = True
    include_clusters: bool = True
    include_embeddings: bool = False
    batch_size: Optional[int] = 1000

class ExportJob(BaseModel):
    job_id: str
    session_id: str
    format: ExportFormat
    status: ExportStatus
    progress: float = 0.0
    total_items: int = 0
    processed_items: int = 0
    file_path: Optional[str] = None
    error_message: Optional[str] = None
    created_at: datetime
    completed_at: Optional[datetime] = None

# In-memory job storage (in production, use Redis or database)
export_jobs: Dict[str, ExportJob] = {}

class ExportEngine:
    """Core export engine for handling different formats."""
    
    def __init__(self):
        self.jinja_env = Environment(
            loader=FileSystemLoader(str(TEMPLATES_DIR)),
            autoescape=True
        )
        self._setup_default_templates()
    
    def _setup_default_templates(self):
        """Create default templates for each export format."""
        default_templates = {
            "markdown_default": """# {{ session_name }} - Export Report

Generated on: {{ export_date }}
Total Items: {{ total_items }}
Total Clusters: {{ total_clusters }}

## Summary
{{ summary }}

{% for cluster in clusters %}
## Cluster {{ cluster.id }}: {{ cluster.label }}

**Size:** {{ cluster.size }} items
**Coherence Score:** {{ cluster.coherence_score }}

### Items in this cluster:
{% for item in cluster['items'] %}
- **{{ item.title }}** ({{ item.domain }})
  - URL: {{ item.url }}
  - Content Preview: {{ item.content[:200] }}...
  
{% endfor %}
{% endfor %}
""",
            "obsidian_default": """---
tags: [export, clustering, {{ session_id }}]
created: {{ export_date }}
total_items: {{ total_items }}
total_clusters: {{ total_clusters }}
---

# {{ session_name }} - Clustering Results

## Overview
- **Session ID:** {{ session_id }}
- **Export Date:** {{ export_date }}
- **Total Items:** {{ total_items }}
- **Total Clusters:** {{ total_clusters }}

## Clusters

{% for cluster in clusters %}
### [[Cluster {{ cluster.id }}]] - {{ cluster.label }}

**Metrics:**
- Size: {{ cluster.size }}
- Coherence: {{ cluster.coherence_score }}

**Items:**
{% for item in cluster['items'] %}
- [[{{ item.title }}]] - {{ item.domain }}
{% endfor %}

{% endfor %}

## Tags
{% for tag in all_tags %}
#{{ tag }} {% endfor %}
""",
            "json_default": """{{ data | tojson(indent=2) }}"""
        }
        
        for template_name, content in default_templates.items():
            template_path = TEMPLATES_DIR / f"{template_name}.j2"
            # Always overwrite to ensure latest template content
            template_path.write_text(content)

    async def get_session_data(self, session_id: str, filters: Optional[ExportFilter] = None) -> Dict[str, Any]:
        """Retrieve and filter session data from Qdrant."""
        try:
            collection_name = f"session_{session_id}"
            
            # Get all points from the collection
            scroll_result = qdrant_client.scroll(
                collection_name=collection_name,
                limit=10000,  # Adjust based on needs
                with_payload=True,
                with_vectors=False
            )
            
            points = scroll_result[0]
            
            # Apply filters if provided
            if filters:
                points = self._apply_filters(points, filters)
            
            # Organize data by clusters
            clusters = {}
            items = []
            
            for point in points:
                payload = point.payload
                cluster_id = payload.get('cluster_id', -1)
                
                item_data = {
                    'id': str(point.id),
                    'title': payload.get('title', 'Untitled'),
                    'url': payload.get('url', ''),
                    'domain': payload.get('domain', ''),
                    'content': payload.get('content', ''),
                    'cluster_id': cluster_id,
                    'cluster_label': payload.get('cluster_label', f'Cluster {cluster_id}'),
                    'created_at': payload.get('created_at', ''),
                    'metadata': payload.get('metadata', {})
                }
                
                items.append(item_data)
                
                if cluster_id not in clusters:
                    clusters[cluster_id] = {
                        'id': cluster_id,
                        'label': payload.get('cluster_label', f'Cluster {cluster_id}'),
                        'items': [],
                        'size': 0,
                        'coherence_score': payload.get('coherence_score', 0.0)
                    }
                
                clusters[cluster_id]['items'].append(item_data)
                clusters[cluster_id]['size'] += 1
            
            return {
                'session_id': session_id,
                'session_name': f'Session {session_id}',
                'items': items,
                'clusters': list(clusters.values()),
                'total_items': len(items),
                'total_clusters': len(clusters),
                'export_date': datetime.now().isoformat(),
                'summary': f'Exported {len(items)} items across {len(clusters)} clusters'
            }
            
        except Exception as e:
            logger.error("Failed to retrieve session data", session_id=session_id, error=str(e))
            raise HTTPException(status_code=500, detail=f"Failed to retrieve session data: {str(e)}")
    
    def _apply_filters(self, points: List[Any], filters: ExportFilter) -> List[Any]:
        """Apply filters to the points data."""
        filtered_points = points
        
        if filters.cluster_ids:
            filtered_points = [p for p in filtered_points 
                             if p.payload.get('cluster_id') in filters.cluster_ids]
        
        if filters.domains:
            filtered_points = [p for p in filtered_points 
                             if p.payload.get('domain') in filters.domains]
        
        if filters.min_score:
            filtered_points = [p for p in filtered_points 
                             if p.payload.get('coherence_score', 0) >= filters.min_score]
        
        if filters.keywords:
            filtered_points = [p for p in filtered_points 
                             if any(keyword.lower() in p.payload.get('content', '').lower() 
                                   for keyword in filters.keywords)]
        
        return filtered_points

    async def export_to_json(self, data: Dict[str, Any], template: Optional[str] = None) -> str:
        """Export data to JSON format."""
        if template:
            template_obj = self.jinja_env.from_string(template)
            return template_obj.render(data=data, **data)
        else:
            return json.dumps(data, indent=2, default=str)

    async def export_to_csv(self, data: Dict[str, Any]) -> str:
        """Export data to CSV format."""
        output = StringIO()
        writer = csv.writer(output)
        
        # Write headers
        headers = ["id", "title", "url", "domain", "cluster_id", "cluster_label", "content_preview", "created_at"]
        writer.writerow(headers)
        
        # Write data rows
        for item in data['items']:
            row = [
                item['id'],
                item['title'],
                item['url'],
                item['domain'],
                item['cluster_id'],
                item['cluster_label'],
                item['content'][:200] + '...' if len(item['content']) > 200 else item['content'],
                item['created_at']
            ]
            writer.writerow(row)
        
        return output.getvalue()

    async def export_to_markdown(self, data: Dict[str, Any], template: Optional[str] = None) -> str:
        """Export data to Markdown format."""
        if template:
            template_obj = self.jinja_env.from_string(template)
        else:
            template_obj = self.jinja_env.get_template("markdown_default.j2")
        
        return template_obj.render(**data)

    async def export_to_obsidian(self, data: Dict[str, Any], template: Optional[str] = None) -> str:
        """Export data to Obsidian-compatible Markdown format."""
        if template:
            template_obj = self.jinja_env.from_string(template)
        else:
            template_obj = self.jinja_env.get_template("obsidian_default.j2")
        
        # Add Obsidian-specific data
        data['all_tags'] = list(set([
            item['domain'].replace('.', '_') for item in data['items']
        ]))
        
        return template_obj.render(**data)

    async def export_to_word(self, data: Dict[str, Any]) -> BytesIO:
        """Export data to Word document format."""
        doc = Document()
        
        # Title
        title = doc.add_heading(f"{data['session_name']} - Export Report", 0)
        
        # Summary
        doc.add_heading('Summary', level=1)
        summary_para = doc.add_paragraph()
        summary_para.add_run(f"Generated on: {data['export_date']}\n")
        summary_para.add_run(f"Total Items: {data['total_items']}\n")
        summary_para.add_run(f"Total Clusters: {data['total_clusters']}\n")
        
        # Clusters
        doc.add_heading('Clusters', level=1)
        
        for cluster in data['clusters']:
            doc.add_heading(f"Cluster {cluster['id']}: {cluster['label']}", level=2)
            
            cluster_info = doc.add_paragraph()
            cluster_info.add_run(f"Size: {cluster['size']} items\n")
            cluster_info.add_run(f"Coherence Score: {cluster['coherence_score']:.3f}\n")
            
            doc.add_heading('Items in this cluster:', level=3)
            
            for item in cluster['items']:
                item_para = doc.add_paragraph(style='List Bullet')
                item_para.add_run(f"{item['title']} ({item['domain']})\n")
                item_para.add_run(f"URL: {item['url']}\n")
                content_preview = item['content'][:200] + '...' if len(item['content']) > 200 else item['content']
                item_para.add_run(f"Content: {content_preview}")
        
        # Save to BytesIO
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        return doc_io

    async def export_to_notion(self, data: Dict[str, Any], notion_token: str, database_id: str) -> Dict[str, Any]:
        """Export data to Notion database."""
        notion = NotionClient(auth=notion_token)
        
        results = []
        
        for item in data['items']:
            try:
                page_data = {
                    "parent": {"database_id": database_id},
                    "properties": {
                        "Title": {
                            "title": [{"text": {"content": item['title']}}]
                        },
                        "URL": {
                            "url": item['url']
                        },
                        "Domain": {
                            "rich_text": [{"text": {"content": item['domain']}}]
                        },
                        "Cluster": {
                            "number": item['cluster_id']
                        },
                        "Cluster Label": {
                            "rich_text": [{"text": {"content": item['cluster_label']}}]
                        }
                    },
                    "children": [
                        {
                            "object": "block",
                            "type": "paragraph",
                            "paragraph": {
                                "rich_text": [{"type": "text", "text": {"content": item['content'][:2000]}}]
                            }
                        }
                    ]
                }
                
                page = notion.pages.create(**page_data)
                results.append({"item_id": item['id'], "notion_page_id": page['id']})
                
            except Exception as e:
                logger.error("Failed to create Notion page", item_id=item['id'], error=str(e))
                results.append({"item_id": item['id'], "error": str(e)})
        
        return {"created_pages": len([r for r in results if 'notion_page_id' in r]), "results": results}

# Initialize export engine
export_engine = ExportEngine()

@app.get("/health")
async def health_check():
    """Health check endpoint."""
    return {
        "status": "healthy",
        "service": "export",
        "timestamp": time.time()
    }

@app.get("/")
async def root():
    """Root endpoint."""
    return {
        "service": "Export Service",
        "version": "1.0.0",
        "status": "running",
        "supported_formats": [format.value for format in ExportFormat]
    }

@app.post("/export", response_model=Dict[str, Any])
async def create_export(
    request: ExportRequest,
    background_tasks: BackgroundTasks
):
    """Create a new export job."""
    job_id = str(uuid.uuid4())
    
    job = ExportJob(
        job_id=job_id,
        session_id=request.session_id,
        format=request.format,
        status=ExportStatus.PENDING,
        created_at=datetime.now()
    )
    
    export_jobs[job_id] = job
    
    # Start background export task
    background_tasks.add_task(process_export_job, job_id, request)
    
    logger.info("Export job created", job_id=job_id, session_id=request.session_id, format=request.format)
    
    return {
        "job_id": job_id,
        "status": job.status,
        "message": "Export job created successfully"
    }

@app.get("/export/{job_id}/status", response_model=ExportJob)
async def get_export_status(job_id: str):
    """Get the status of an export job."""
    if job_id not in export_jobs:
        raise HTTPException(status_code=404, detail="Export job not found")
    
    return export_jobs[job_id]

@app.get("/export/{job_id}/download")
async def download_export(job_id: str):
    """Download the exported file."""
    if job_id not in export_jobs:
        raise HTTPException(status_code=404, detail="Export job not found")
    
    job = export_jobs[job_id]
    
    if job.status != ExportStatus.COMPLETED:
        raise HTTPException(status_code=400, detail="Export job not completed")
    
    if not job.file_path or not Path(job.file_path).exists():
        raise HTTPException(status_code=404, detail="Export file not found")
    
    from fastapi.responses import FileResponse
    
    filename = f"export_{job.session_id}_{job.format.value}_{job_id[:8]}"
    
    return FileResponse(
        path=job.file_path,
        filename=filename,
        media_type="application/octet-stream"
    )

@app.post("/export/batch", response_model=Dict[str, Any])
async def create_batch_export(
    requests: List[ExportRequest],
    background_tasks: BackgroundTasks
):
    """Create multiple export jobs for batch processing."""
    job_ids = []
    
    for request in requests:
        job_id = str(uuid.uuid4())
        
        job = ExportJob(
            job_id=job_id,
            session_id=request.session_id,
            format=request.format,
            status=ExportStatus.PENDING,
            created_at=datetime.now()
        )
        
        export_jobs[job_id] = job
        job_ids.append(job_id)
        
        # Start background export task
        background_tasks.add_task(process_export_job, job_id, request)
    
    logger.info("Batch export jobs created", job_count=len(job_ids))
    
    return {
        "job_ids": job_ids,
        "total_jobs": len(job_ids),
        "message": "Batch export jobs created successfully"
    }

@app.get("/export/jobs", response_model=List[ExportJob])
async def list_export_jobs(
    session_id: Optional[str] = Query(None),
    status: Optional[ExportStatus] = Query(None),
    limit: int = Query(100, ge=1, le=1000)
):
    """List export jobs with optional filtering."""
    jobs = list(export_jobs.values())
    
    if session_id:
        jobs = [job for job in jobs if job.session_id == session_id]
    
    if status:
        jobs = [job for job in jobs if job.status == status]
    
    # Sort by creation date (newest first)
    jobs.sort(key=lambda x: x.created_at, reverse=True)
    
    return jobs[:limit]

@app.delete("/export/{job_id}")
async def delete_export_job(job_id: str):
    """Delete an export job and its associated file."""
    if job_id not in export_jobs:
        raise HTTPException(status_code=404, detail="Export job not found")
    
    job = export_jobs[job_id]
    
    # Delete file if it exists
    if job.file_path and Path(job.file_path).exists():
        try:
            Path(job.file_path).unlink()
        except Exception as e:
            logger.warning("Failed to delete export file", file_path=job.file_path, error=str(e))
    
    # Remove job from memory
    del export_jobs[job_id]
    
    logger.info("Export job deleted", job_id=job_id)
    
    return {"message": "Export job deleted successfully"}

@app.post("/templates", response_model=Dict[str, Any])
async def create_template(template: ExportTemplate):
    """Create a custom export template."""
    template_path = TEMPLATES_DIR / f"{template.name}_{template.format.value}.j2"
    
    if template_path.exists():
        raise HTTPException(status_code=400, detail="Template already exists")
    
    try:
        template_path.write_text(template.template_content)
        logger.info("Template created", name=template.name, format=template.format)
        
        return {
            "message": "Template created successfully",
            "template_name": template.name,
            "template_path": str(template_path)
        }
    except Exception as e:
        logger.error("Failed to create template", name=template.name, error=str(e))
        raise HTTPException(status_code=500, detail=f"Failed to create template: {str(e)}")

@app.get("/templates", response_model=List[Dict[str, Any]])
async def list_templates():
    """List available export templates."""
    templates = []
    
    for template_file in TEMPLATES_DIR.glob("*.j2"):
        name_parts = template_file.stem.split("_")
        if len(name_parts) >= 2:
            name = "_".join(name_parts[:-1])
            format_name = name_parts[-1]
        else:
            name = template_file.stem
            format_name = "unknown"
        
        templates.append({
            "name": name,
            "format": format_name,
            "file_path": str(template_file),
            "created_at": datetime.fromtimestamp(template_file.stat().st_mtime).isoformat()
        })
    
    return templates

@app.get("/templates/{template_name}")
async def get_template(template_name: str, format: ExportFormat):
    """Get a specific template content."""
    template_path = TEMPLATES_DIR / f"{template_name}_{format.value}.j2"
    
    if not template_path.exists():
        raise HTTPException(status_code=404, detail="Template not found")
    
    try:
        content = template_path.read_text()
        return {
            "name": template_name,
            "format": format.value,
            "content": content,
            "file_path": str(template_path)
        }
    except Exception as e:
        logger.error("Failed to read template", template_name=template_name, error=str(e))
        raise HTTPException(status_code=500, detail=f"Failed to read template: {str(e)}")

async def process_export_job(job_id: str, request: ExportRequest):
    """Background task to process export jobs."""
    job = export_jobs[job_id]
    
    try:
        job.status = ExportStatus.PROCESSING
        logger.info("Starting export job", job_id=job_id, format=request.format)
        
        # Get session data
        data = await export_engine.get_session_data(request.session_id, request.filters)
        job.total_items = data['total_items']
        
        # Determine template
        template_content = None
        if request.custom_template:
            template_content = request.custom_template
        elif request.template_name:
            template_path = TEMPLATES_DIR / f"{request.template_name}_{request.format.value}.j2"
            if template_path.exists():
                template_content = template_path.read_text()
        
        # Process export based on format
        file_extension = request.format.value
        output_filename = f"export_{request.session_id}_{request.format.value}_{job_id[:8]}.{file_extension}"
        output_path = EXPORT_DIR / output_filename
        
        if request.format == ExportFormat.JSON:
            content = await export_engine.export_to_json(data, template_content)
            output_path.write_text(content)
            
        elif request.format == ExportFormat.CSV:
            content = await export_engine.export_to_csv(data)
            output_path.write_text(content)
            
        elif request.format == ExportFormat.MARKDOWN:
            content = await export_engine.export_to_markdown(data, template_content)
            output_path.write_text(content)
            
        elif request.format == ExportFormat.OBSIDIAN:
            content = await export_engine.export_to_obsidian(data, template_content)
            output_path = output_path.with_suffix('.md')
            output_path.write_text(content)
            
        elif request.format == ExportFormat.WORD:
            doc_io = await export_engine.export_to_word(data)
            output_path = output_path.with_suffix('.docx')
            output_path.write_bytes(doc_io.getvalue())
            
        elif request.format == ExportFormat.NOTION:
            # For Notion, we need additional parameters
            # This would typically come from the request
            raise HTTPException(status_code=400, detail="Notion export requires additional configuration")
            
        else:
            raise ValueError(f"Unsupported export format: {request.format}")
        
        # Update job status
        job.status = ExportStatus.COMPLETED
        job.processed_items = job.total_items
        job.progress = 100.0
        job.file_path = str(output_path)
        job.completed_at = datetime.now()
        
        logger.info("Export job completed", job_id=job_id, file_path=str(output_path))
        
    except Exception as e:
        job.status = ExportStatus.FAILED
        job.error_message = str(e)
        logger.error("Export job failed", job_id=job_id, error=str(e))

@app.get("/formats", response_model=List[Dict[str, Any]])
async def get_supported_formats():
    """Get list of supported export formats with descriptions."""
    formats = [
        {
            "format": ExportFormat.JSON.value,
            "description": "JavaScript Object Notation - structured data format",
            "file_extension": "json",
            "supports_templates": True
        },
        {
            "format": ExportFormat.CSV.value,
            "description": "Comma-Separated Values - spreadsheet compatible format",
            "file_extension": "csv",
            "supports_templates": False
        },
        {
            "format": ExportFormat.MARKDOWN.value,
            "description": "Markdown format - human-readable text format",
            "file_extension": "md",
            "supports_templates": True
        },
        {
            "format": ExportFormat.OBSIDIAN.value,
            "description": "Obsidian-compatible Markdown with internal linking",
            "file_extension": "md",
            "supports_templates": True
        },
        {
            "format": ExportFormat.WORD.value,
            "description": "Microsoft Word document format",
            "file_extension": "docx",
            "supports_templates": False
        },
        {
            "format": ExportFormat.NOTION.value,
            "description": "Export to Notion database",
            "file_extension": None,
            "supports_templates": False
        }
    ]
    
    return formats

@app.post("/export/preview", response_model=Dict[str, Any])
async def preview_export(
    session_id: str,
    format: ExportFormat,
    template_content: Optional[str] = None,
    filters: Optional[ExportFilter] = None,
    limit: int = Query(5, ge=1, le=50)
):
    """Preview export output with limited data."""
    try:
        # Get limited session data
        data = await export_engine.get_session_data(session_id, filters)
        
        # Limit data for preview
        data['items'] = data['items'][:limit]
        data['clusters'] = [
            {**cluster, 'items': cluster['items'][:limit]}
            for cluster in data['clusters'][:3]  # Show max 3 clusters
        ]
        data['total_items'] = len(data['items'])
        data['total_clusters'] = len(data['clusters'])
        
        # Generate preview based on format
        if format == ExportFormat.JSON:
            preview = await export_engine.export_to_json(data, template_content)
        elif format == ExportFormat.CSV:
            preview = await export_engine.export_to_csv(data)
        elif format == ExportFormat.MARKDOWN:
            preview = await export_engine.export_to_markdown(data, template_content)
        elif format == ExportFormat.OBSIDIAN:
            preview = await export_engine.export_to_obsidian(data, template_content)
        else:
            preview = "Preview not available for this format"
        
        return {
            "preview": preview[:2000],  # Limit preview size
            "truncated": len(preview) > 2000,
            "sample_items": len(data['items']),
            "sample_clusters": len(data['clusters'])
        }
        
    except Exception as e:
        logger.error("Failed to generate preview", session_id=session_id, error=str(e))
        raise HTTPException(status_code=500, detail=f"Failed to generate preview: {str(e)}")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)