"""Core export engine implementation."""

from __future__ import annotations

import csv
import json
from datetime import datetime
from io import BytesIO, StringIO
from pathlib import Path
from typing import Any, Callable, Dict, List, Optional

try:  # Optional dependency. See README for installation guidance.
    from docx import Document  # type: ignore
except ImportError:  # pragma: no cover - handled at runtime
    Document = None  # type: ignore
from fastapi import HTTPException
from jinja2 import Environment, FileSystemLoader
from notion_client import Client as NotionClient
from qdrant_client import QdrantClient
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from .models import ExportFilter
from .templates import ensure_default_templates


class ExportEngine:
    """Core export engine responsible for data retrieval and formatting."""

    def __init__(
        self,
        qdrant_client: QdrantClient,
        templates_dir: Path,
        notion_client_factory: Optional[Callable[..., Any]] = None,
    ) -> None:
        self.qdrant_client = qdrant_client
        self.templates_dir = templates_dir
        self.notion_client_factory = notion_client_factory or NotionClient

        ensure_default_templates(self.templates_dir)
        self.jinja_env = Environment(
            loader=FileSystemLoader(str(self.templates_dir)),
            autoescape=True,
        )

    async def get_session_data(
        self,
        session_id: str,
        filters: Optional[ExportFilter] = None,
    ) -> Dict[str, Any]:
        """Retrieve session data from Qdrant, optionally applying filters."""
        try:
            collection_name = f"session_{session_id}"
            scroll_result = self.qdrant_client.scroll(
                collection_name=collection_name,
                limit=10000,
                with_payload=True,
                with_vectors=False,
            )
            points = scroll_result[0]

            if filters:
                points = self._apply_filters(points, filters)

            clusters: Dict[int, Dict[str, Any]] = {}
            items: List[Dict[str, Any]] = []

            for point in points:
                payload = point.payload
                cluster_id = payload.get("cluster_id", -1)

                item_data = {
                    "id": str(point.id),
                    "title": payload.get("title", "Untitled"),
                    "url": payload.get("url", ""),
                    "domain": payload.get("domain", ""),
                    "content": payload.get("content", ""),
                    "cluster_id": cluster_id,
                    "cluster_label": payload.get("cluster_label", f"Cluster {cluster_id}"),
                    "created_at": payload.get("created_at", ""),
                    "metadata": payload.get("metadata", {}),
                }

                items.append(item_data)

                if cluster_id not in clusters:
                    clusters[cluster_id] = {
                        "id": cluster_id,
                        "label": payload.get("cluster_label", f"Cluster {cluster_id}"),
                        "items": [],
                        "size": 0,
                        "coherence_score": payload.get("coherence_score", 0.0),
                    }

                clusters[cluster_id]["items"].append(item_data)
                clusters[cluster_id]["size"] += 1

            return {
                "session_id": session_id,
                "session_name": f"Session {session_id}",
                "items": items,
                "clusters": list(clusters.values()),
                "total_items": len(items),
                "total_clusters": len(clusters),
                "export_date": datetime.now().isoformat(),
                "summary": f"Exported {len(items)} items across {len(clusters)} clusters",
            }
        except Exception as exc:  # pragma: no cover - defensive logging
            raise HTTPException(status_code=500, detail=f"Failed to retrieve session data: {exc}")

    def _apply_filters(self, points: List[Any], filters: ExportFilter) -> List[Any]:
        """Apply filtering logic to raw Qdrant points."""
        filtered_points = points

        if filters.cluster_ids:
            filtered_points = [
                point for point in filtered_points if point.payload.get("cluster_id") in filters.cluster_ids
            ]

        if filters.domains:
            filtered_points = [
                point for point in filtered_points if point.payload.get("domain") in filters.domains
            ]

        if filters.min_score:
            filtered_points = [
                point
                for point in filtered_points
                if point.payload.get("coherence_score", 0) >= filters.min_score
            ]

        if filters.keywords:
            filtered_points = [
                point
                for point in filtered_points
                if any(
                    keyword.lower() in point.payload.get("content", "").lower()
                    for keyword in filters.keywords
                )
            ]

        return filtered_points

    async def export_to_json(self, data: Dict[str, Any], template: Optional[str] = None) -> str:
        """Render data into JSON, optionally using a custom template."""
        if template:
            template_obj = self.jinja_env.from_string(template)
            return template_obj.render(data=data, **data)
        return json.dumps(data, indent=2, default=str)

    async def export_to_csv(self, data: Dict[str, Any]) -> str:
        """Render data into CSV format."""
        output = StringIO()
        writer = csv.writer(output)
        headers = [
            "id",
            "title",
            "url",
            "domain",
            "cluster_id",
            "cluster_label",
            "content_preview",
            "created_at",
        ]
        writer.writerow(headers)

        for item in data["items"]:
            row = [
                item["id"],
                item["title"],
                item["url"],
                item["domain"],
                item["cluster_id"],
                item["cluster_label"],
                item["content"][:200] + "..." if len(item["content"]) > 200 else item["content"],
                item["created_at"],
            ]
            writer.writerow(row)

        return output.getvalue()

    async def export_to_markdown(self, data: Dict[str, Any], template: Optional[str] = None) -> str:
        """Render data into Markdown format."""
        template_obj = (
            self.jinja_env.from_string(template)
            if template
            else self.jinja_env.get_template("markdown_default.j2")
        )
        return template_obj.render(**data)

    async def export_to_obsidian(self, data: Dict[str, Any], template: Optional[str] = None) -> str:
        """Render data into Obsidian markdown, enriching with tags."""
        template_obj = (
            self.jinja_env.from_string(template)
            if template
            else self.jinja_env.get_template("obsidian_default.j2")
        )

        render_data = dict(data)
        render_data["all_tags"] = list({item["domain"].replace(".", "_") for item in data["items"]})
        return template_obj.render(**render_data)

    async def export_to_word(self, data: Dict[str, Any]) -> BytesIO:
        """Create a Word document containing the export data."""
        if Document is None:
            raise RuntimeError("python-docx is not installed; Word export is unavailable.")

        document = Document()
        document.add_heading(f"{data['session_name']} - Export Report", 0)
        document.add_heading("Summary", level=1)

        summary_para = document.add_paragraph()
        summary_para.add_run(f"Generated on: {data['export_date']}\n")
        summary_para.add_run(f"Total Items: {data['total_items']}\n")
        summary_para.add_run(f"Total Clusters: {data['total_clusters']}\n")

        document.add_heading("Clusters", level=1)

        for cluster in data["clusters"]:
            document.add_heading(f"Cluster {cluster['id']}: {cluster['label']}", level=2)

            cluster_info = document.add_paragraph()
            cluster_info.add_run(f"Size: {cluster['size']} items\n")
            cluster_info.add_run(f"Coherence Score: {cluster['coherence_score']:.3f}\n")

            document.add_heading("Items in this cluster:", level=3)

            for item in cluster["items"]:
                item_para = document.add_paragraph(style="List Bullet")
                item_para.add_run(f"{item['title']} ({item['domain']})\n")
                item_para.add_run(f"URL: {item['url']}\n")
                content_preview = (
                    item["content"][:200] + "..." if len(item["content"]) > 200 else item["content"]
                )
                item_para.add_run(f"Content: {content_preview}")

        doc_io = BytesIO()
        document.save(doc_io)
        doc_io.seek(0)
        return doc_io

    async def export_to_notion(
        self,
        data: Dict[str, Any],
        notion_token: str,
        database_id: str,
    ) -> Dict[str, Any]:
        """Push export data into a Notion database."""
        notion = self.notion_client_factory(auth=notion_token)
        results = []

        for item in data["items"]:
            try:
                page_data = {
                    "parent": {"database_id": database_id},
                    "properties": {
                        "Title": {"title": [{"text": {"content": item["title"]}}]},
                        "URL": {"url": item["url"]},
                        "Domain": {"rich_text": [{"text": {"content": item["domain"]}}]},
                        "Cluster": {"number": item["cluster_id"]},
                        "Cluster Label": {
                            "rich_text": [{"text": {"content": item["cluster_label"]}}]
                        },
                    },
                    "children": [
                        {
                            "object": "block",
                            "type": "paragraph",
                            "paragraph": {
                                "rich_text": [
                                    {"type": "text", "text": {"content": item["content"][:2000]}}
                                ]
                            },
                        }
                    ],
                }

                page = notion.pages.create(**page_data)
                results.append({"item_id": item["id"], "notion_page_id": page["id"]})
            except Exception as exc:  # pragma: no cover - logging retained in caller
                results.append({"item_id": item["id"], "error": str(exc)})

        created_pages = len([result for result in results if "notion_page_id" in result])
        return {"created_pages": created_pages, "results": results}

    async def export_to_pdf(self, data: Dict[str, Any]) -> BytesIO:
        """Create a PDF document summarising the export data."""
        pdf_io = BytesIO()
        document = SimpleDocTemplate(
            pdf_io,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18,
        )

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "CustomTitle",
            parent=styles["Heading1"],
            fontSize=24,
            spaceAfter=30,
            alignment=1,
        )
        heading_style = ParagraphStyle(
            "CustomHeading",
            parent=styles["Heading2"],
            fontSize=16,
            spaceAfter=12,
            textColor=colors.darkblue,
        )
        normal_style = styles["Normal"]

        story = []
        story.append(Paragraph(f"{data['session_name']} - Export Report", title_style))
        story.append(Spacer(1, 20))

        summary_data = [
            ["Generated on:", data["export_date"]],
            ["Total Items:", str(data["total_items"])],
            ["Total Clusters:", str(data["total_clusters"])],
            ["Summary:", data["summary"]],
        ]

        summary_table = Table(summary_data, colWidths=[2 * inch, 4 * inch])
        summary_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (0, -1), colors.lightgrey),
                    ("TEXTCOLOR", (0, 0), (-1, -1), colors.black),
                    ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                    ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
                    ("FONTSIZE", (0, 0), (-1, -1), 10),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 12),
                    ("BACKGROUND", (1, 0), (1, -1), colors.beige),
                    ("GRID", (0, 0), (-1, -1), 1, colors.black),
                ]
            )
        )
        story.append(summary_table)
        story.append(Spacer(1, 30))

        story.append(Paragraph("Clusters", heading_style))
        story.append(Spacer(1, 12))

        for cluster in data["clusters"]:
            story.append(Paragraph(f"Cluster {cluster['id']}: {cluster['label']}", styles["Heading3"]))

            cluster_info = [
                ["Size:", f"{cluster['size']} items"],
                ["Coherence Score:", f"{cluster['coherence_score']:.3f}"],
            ]
            cluster_table = Table(cluster_info, colWidths=[1.5 * inch, 2 * inch])
            cluster_table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (0, -1), colors.lightblue),
                        ("TEXTCOLOR", (0, 0), (-1, -1), colors.black),
                        ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                        ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
                        ("FONTSIZE", (0, 0), (-1, -1), 9),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                        ("GRID", (0, 0), (-1, -1), 1, colors.black),
                    ]
                )
            )
            story.append(cluster_table)
            story.append(Spacer(1, 12))

            story.append(Paragraph("Items in this cluster:", styles["Heading4"]))

            items_data = [["Title", "Domain", "URL"]]
            for item in cluster["items"][:10]:
                title = item["title"][:50] + "..." if len(item["title"]) > 50 else item["title"]
                domain = item["domain"]
                url = item["url"][:60] + "..." if len(item["url"]) > 60 else item["url"]
                items_data.append([title, domain, url])

            if len(cluster["items"]) > 10:
                items_data.append([f"... and {len(cluster['items']) - 10} more items", "", ""])

            items_table = Table(items_data, colWidths=[2.5 * inch, 1.5 * inch, 2.5 * inch])
            items_table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                        ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                        ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                        ("FONTSIZE", (0, 0), (-1, -1), 8),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                        ("GRID", (0, 0), (-1, -1), 1, colors.black),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ]
                )
            )
            story.append(items_table)
            story.append(Spacer(1, 20))

        document.build(story)
        pdf_io.seek(0)
        return pdf_io


__all__ = ["ExportEngine"]
